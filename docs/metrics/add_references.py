"""Appends a References chapter -- the manuscript had none at all, not even
for the RRL's own six studies. Every entry here was verified against a real
source (web search + the Roboflow exports' own bundled README/data.yaml
files), not reconstructed from memory. APA 7th edition format, matching what
the rest of the manuscript's in-text citations (Author, Year) assume.

Run from repo root:
    .venv\\Scripts\\python.exe docs\\metrics\\add_references.py
"""
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = r"D:\projects\EcoVisionCode\docs\EcoVision - Manuscript (REVISED).docx"
d = docx.Document(PATH)

if any(p.text.strip() == "REFERENCES" for p in d.paragraphs):
    raise SystemExit("REFERENCES already present -- refusing to duplicate.")

# ---------------------------------------------------------------------------
# Heading, matching "LIST OF FIGURES" style exactly (No Spacing, bold, centred)
# ---------------------------------------------------------------------------
h = d.add_paragraph(style="No Spacing")
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h.add_run("REFERENCES")
r.font.bold = True
r.font.size = Pt(12)

intro = d.add_paragraph()
intro.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro.add_run(
    "All in-text citations in this manuscript, and every dataset used in Chapter III, are "
    "listed below in APA 7th edition format. Entries were verified against their original "
    "source rather than reconstructed from memory."
)

# ---------------------------------------------------------------------------
# References, alphabetical by first author/organisation surname
# ---------------------------------------------------------------------------
refs = [
    "Ahmed, M. (2022). gun and knife detection (Version 1) [Data set]. Roboflow. "
    "https://universe.roboflow.com/mahad-ahmed/gun-and-knife-detection",

    "Amah, G. G., Enokela, J. A., Agbo, D. O., & Iorkyase, T. E. (2025). Lightweight YOLOv8 "
    "optimized deep neural network for real-time weapon detection on Raspberry Pi 5 in smart "
    "surveillance systems. Journal of Engineering Research and Reports, 27(9). "
    "https://journaljerr.com/index.php/JERR/article/view/1688",

    "Aremu, T., Zhiyuan, L., Alameeri, R., Khan, M., & El Saddik, A. (2024). SSIVD-Net: A "
    "novel salient super image classification & detection technique for weaponized violence "
    "[Introduces the Smart-City CCTV Violence Detection (SCVD) dataset]. Proceedings of the "
    "2024 SAI Computing Conference. https://arxiv.org/abs/2207.12850",

    "Berg Insight. (2025). The global smart street lighting market (3rd ed.) [Market research "
    "report]. https://www.berginsight.com/the-global-smart-street-lighting-market/",

    "Bosun Lighting. (2025). GEBOSUN lighting and smart cities: Smart light poles [Company "
    "product documentation]. "
    "https://www.bosunlighting.com/gebosun-lighting-and-smart-cities-smart-light-poles.html",

    "Chaudhry, U. (2022). Traffic and Road Signs (Version 1) [Data set]. Roboflow. "
    "https://universe.roboflow.com/usmanchaudhry622-gmail-com/traffic-and-road-signs",

    "Cheng, M., Cai, K., & Li, M. (2021). RWF-2000: An open large scale video database for "
    "violence detection. In 2020 25th International Conference on Pattern Recognition (ICPR) "
    "(pp. 4183-4190). IEEE. https://doi.org/10.1109/ICPR48806.2021.9412502",

    "dietest. (2023). Gun-cctv-detection (Version 1) [Data set]. Roboflow. "
    "https://universe.roboflow.com/dietest/gun-cctv-detection",

    "Florin, A. F., Ugalino, M., Jr., & Aguilan, K. E. (2025). An analysis of smart city "
    "development frameworks [Policy brief]. UP Center for Integrative and Development "
    "Studies. https://cids.up.edu.ph/policy-brief/analysis-smart-city-development-frameworks/",

    "Gao, H. (2023). A YOLO-based violence detection method in IoT surveillance systems. "
    "International Journal of Advanced Computer Science and Applications, 14(8). "
    "https://doi.org/10.14569/IJACSA.2023.0140817",

    "Kulkarni, S. (2022). knife (Version 1) [Data set]. Roboflow. "
    "https://universe.roboflow.com/sanket-kulkarni/knife-eydvx",

    "Perez, M., Kot, A. C., & Rocha, A. (2019). Detection of real-world fights in surveillance "
    "videos [Introduces the CCTV-Fights dataset]. ICASSP 2019 - 2019 IEEE International "
    "Conference on Acoustics, Speech and Signal Processing (ICASSP), 2662-2666. "
    "https://doi.org/10.1109/ICASSP.2019.8683676",

    "Pudasaini, D., & Abhari, A. (2021). Edge-based video analytic for smart cities. "
    "International Journal of Advanced Computer Science and Applications, 12(7). "
    "https://doi.org/10.14569/IJACSA.2021.0120701",

    "Republic of the Philippines. (2001). Republic Act No. 9003: Ecological Solid Waste "
    "Management Act of 2000. Official Gazette of the Republic of the Philippines. "
    "https://www.officialgazette.gov.ph/2001/01/26/republic-act-no-9003-s-2001/",

    "Simuletic. (2025). CCTV Knife Detection Dataset (Sample Version) (Version 1) [Data set]. "
    "Roboflow. https://universe.roboflow.com/simuletic/cctv-knife-detection-dataset-zkkaf",

    "Sultani, W., Chen, C., & Shah, M. (2018). Real-world anomaly detection in surveillance "
    "videos [Introduces the UCF-Crime dataset]. Proceedings of the IEEE Conference on Computer "
    "Vision and Pattern Recognition (CVPR), 6479-6488. "
    "https://doi.org/10.1109/CVPR.2018.00678",

    "Uranus, H., Adhinugroho, N. R., Yulian, D. H., & Mangunsong, R. (2022). Design and "
    "realization of solar-powered IoT-based flood early warning system with Telegram "
    "messaging, auto-restart watchdog, and power management. GCISTEM Proceeding, 1. "
    "https://doi.org/10.56573/gcistem.v1i.4",

    "workspace-1qko2. (2024). gun detection (Version 4) [Data set]. Roboflow. "
    "https://universe.roboflow.com/workspace-1qko2/gun-detection-ghlzd",

    "workspace-zqssx. (2022). knife-dataset (Version 2) [Data set]. Roboflow. "
    "https://universe.roboflow.com/workspace-zqssx/knife-dataset-4kytl",
]

for entry in refs:
    p = d.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # APA hanging indent: 0.5" left indent, -0.5" first-line (i.e. first line
    # sits at the margin, every wrapped line indents 0.5").
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(entry)
    r.font.size = Pt(11)

note = d.add_paragraph()
note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
note.paragraph_format.space_before = Pt(12)
r = note.add_run(
    "Note: workspace-1qko2, workspace-zqssx, and dietest are Roboflow workspace handles; no "
    "further author name was published on the source pages. Bosun Lighting's entry is a "
    "company product page, not a peer-reviewed source, cited here because Chapter II relies "
    "on it as an industry case study rather than as a research finding -- treat its claims "
    "with correspondingly less weight than the peer-reviewed entries above it."
)
r.font.italic = True
r.font.size = Pt(9)

d.save(PATH)
print(f"Added REFERENCES chapter with {len(refs)} entries.")
