"""Appends the AI findings section (4.11-4.13) to the manuscript.

Run from repo root:
    .venv\\Scripts\\python.exe docs\\metrics\\insert_ai_findings.py

Idempotent-ish: refuses to run twice (checks for the 4.11 heading already
present) so re-running after a manual edit doesn't duplicate the section.
"""
import json
from pathlib import Path

import docx
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

REPO = Path(r"D:\projects\EcoVisionCode")
DOCX_PATH = REPO / "docs" / "EcoVision - Manuscript (REVISED).docx"
FIG = REPO / "docs" / "metrics" / "figures"
RESULTS = json.loads((REPO / "docs" / "metrics" / "confusion_matrix_results.json").read_text())

d = docx.Document(str(DOCX_PATH))

# Guard against double-insertion.
if any("4.11 Confusion Matrix Results" in p.text for p in d.paragraphs):
    raise SystemExit("4.11 already present -- refusing to insert a duplicate. "
                      "Delete the existing 4.11-4.13 section first if you want to regenerate it.")


def add_heading(text):
    p = d.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text)
    r.font.bold = True
    return p


def add_body(text):
    p = d.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def add_caption(text):
    p = d.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(10)
    return p


def add_figure(path, width_in):
    p = d.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    return p


def set_cell_text(cell, text, bold=False, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.paragraph_format.alignment = align
    r = p.add_run(text)
    r.font.bold = bold
    r.font.size = Pt(size)


# ===========================================================================
# 4.11 Confusion Matrix Results
# ===========================================================================
add_heading("4.11 Confusion Matrix Results for All Three Detectors")
add_body(
    "Sections 4.1 through 4.10 report accuracy, recall and false-positive rate as summary "
    "percentages. The raw true-positive, false-positive, true-negative and false-negative counts "
    "behind those percentages are given here in full, evaluated directly against the same code path "
    "production runs (SceneViolenceDetector in maincode/x3d_violence_detector.py), one prediction per "
    "clip, at the exact threshold each detector is currently configured with in config.json. All figures are measured "
    "on each detector's held-out TEST split -- clips excluded from both training and checkpoint "
    "selection for that detector's manifest."
)

tbl = d.add_table(rows=4, cols=8)
tbl.style = "Table Grid"
headers = ["Detector", "Threshold", "TP", "FP", "TN", "FN", "N", "Accuracy"]
for j, h in enumerate(headers):
    set_cell_text(tbl.rows[0].cells[j], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for i, spec in enumerate(RESULTS, start=1):
    name = spec["name"].split(" (")[0]
    row = tbl.rows[i]
    vals = [name, str(spec["threshold"]), str(spec["tp"]), str(spec["fp"]),
            str(spec["tn"]), str(spec["fn"]), str(spec["n"]), f"{spec['accuracy']*100:.1f}%"]
    for j, v in enumerate(vals):
        set_cell_text(row.cells[j], v, align=WD_ALIGN_PARAGRAPH.CENTER)
add_caption("Table 4.6: Confusion Matrix Counts at Configured Operating Points")

# three confusion grids side by side via a borderless table
img_tbl = d.add_table(rows=1, cols=3)
img_tbl.autofit = True
safe_names = ["confusion_violence__daynight_checkpoint.png",
              "confusion_robbery.png",
              "confusion_vandalism__currently_disabled_in_config.png"]
for j, fname in enumerate(safe_names):
    cell = img_tbl.rows[0].cells[j]
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIG / fname), width=Inches(2.05))
# strip borders from this layout-only table
tblPr = img_tbl._tbl.tblPr
borders = tblPr.makeelement(qn("w:tblBorders"), {})
for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
    el = borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "none"})
    borders.append(el)
tblPr.append(borders)
add_caption("Figure 4.11a-c: Confusion Matrices — Violence (daynight checkpoint), Robbery, Vandalism")

add_figure(FIG / "summary_bars.png", 6.0)
add_caption("Figure 4.11d: Recall, Precision and False-Positive Rate by Detector")

add_body(
    "Violence detection is the most mature of the three: 90.1 percent accuracy with a 3.7 percent "
    "false-positive rate on 1,310 test clips, the model discussed throughout this chapter. Robbery "
    "reaches 86.5 percent precision but only 65.3 percent recall on a much smaller test set of 139 "
    "clips, consistent with training on the smaller labelled corpus available for that category "
    "(Chapter III, Dataset Development). Vandalism's figures -- 75.9 percent recall but a 37.5 percent "
    "false-positive rate -- should be read with the same caution given to any measurement on 37 test "
    "clips: a single misclassified clip moves the reported rate by roughly 2.7 percentage points, so "
    "this is a noisy estimate rather than a settled result. This is precisely why config.json currently "
    "ships vandalism detection disabled (detection.vandalism.enabled = false) rather than presenting it "
    "as validated to the same standard as violence and robbery -- an explicit scope decision recorded in "
    "Section 4.10 rather than an oversight. None of the three figures reported in this section are "
    "field-validated: all were measured on held-out benchmark and evaluation footage, not on incidents "
    "observed through an installed, operating system."
)

# ===========================================================================
# 4.12 Detection Architecture
# ===========================================================================
add_heading("4.12 Detection Architecture: The Per-Frame Pipeline")
add_body(
    "Every camera frame passes through the same fixed pipeline. YOLO11s-pose first locates people and "
    "body keypoints in the frame. Rather than three unrelated models, violence, robbery and vandalism "
    "share one detector class (SceneViolenceDetector) parameterised by a different X3D-XS checkpoint, "
    "confidence threshold and consecutive-frame requirement per category, read from config.json. Each "
    "detector's raw per-frame confidence is passed through _smooth_and_confirm, which applies an "
    "exponential moving average and requires a configurable number of consecutive above-threshold frames "
    "before an incident is confirmed -- a single high-confidence frame does not by itself raise an alert. "
    "Weapon and sign detection runs as a separate YOLO model outside this shared class, since it is a "
    "single-frame object-detection task rather than a temporal one."
)
add_figure(FIG / "architecture_pipeline.png", 6.3)
add_caption("Figure 4.12: Per-Frame Detection Pipeline for One Camera")
add_body(
    "X3D-XS was chosen specifically because it is small enough (approximately 3 million parameters) to "
    "keep three instances resident on a single consumer GPU (the development machine's GTX 1660 SUPER, "
    "6 GB VRAM) with headroom left for the pose and weapon models running alongside them, as confirmed by "
    "preflight.py's inference benchmark (Chapter III, Development Phase)."
)

# ===========================================================================
# 4.13 Concurrent Multi-Detector Execution
# ===========================================================================
add_heading("4.13 Concurrent Multi-Detector Execution")
add_body(
    "\"Running at the same time\" in this system means two distinct things, both visible in "
    "maincode/main.py. First, the three category detectors are not three separate processes racing each "
    "other -- they run sequentially, in order, on every frame within a single main detection loop, which "
    "is affordable specifically because each X3D-XS pass is inexpensive. Second, and separately, four "
    "supporting tasks -- weapon/sign detection, posting a confirmed alert to the backend, finalising a "
    "saved incident clip, and encoding frames for the live dashboard stream -- are each offloaded to "
    "their own dedicated single-worker thread pool (Python's ThreadPoolExecutor) rather than run inline. "
    "A camera frame is read on its own thread as well, so a slow network write or a slow encode never "
    "stalls frame capture or the next detection pass."
)
add_figure(FIG / "concurrency_model.png", 6.3)
add_caption("Figure 4.13: Concurrency Model Within One Camera Process")
add_body(
    "This design keeps one camera responsive under load, but it does not by itself scale to many "
    "cameras: the current implementation runs one such pipeline per camera process, with no inference "
    "shared across cameras. Person-crop inference, which would run the classifiers only where people are "
    "detected rather than on every frame regardless of activity, is recorded in Section 4.10 as the "
    "principal remaining step toward serving multiple cameras from one GPU."
)

d.save(str(DOCX_PATH))
print("Inserted sections 4.11-4.13. Saved:", DOCX_PATH)
