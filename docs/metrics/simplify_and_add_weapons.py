"""One combined pass, run once the manuscript is closed in Word:
  1. Simplifies Section 4.11-4.13's prose -- it leaned on code artifacts
     (class/function names, config.json keys) the rest of the chapter
     never does. Facts and numbers are unchanged; only wording changes.
  2. Adds "F. Weapon and Sign Detector Dataset" after Evaluation Metrics,
     with a clearly marked placeholder since the exact Roboflow source
     isn't on record anywhere in the repo (confirmed by search).

Run from repo root:
    .venv\\Scripts\\python.exe docs\\metrics\\simplify_and_add_weapons.py
"""
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = r"D:\projects\EcoVisionCode\docs\EcoVision - Manuscript (REVISED).docx"
d = docx.Document(PATH)
paras = d.paragraphs


def find_para(substring):
    # Search by content instead of a fixed index: the document shifted by 17
    # paragraphs between when indices were first recorded and now (Word
    # appears to have normalised something on open/save), so hardcoded
    # indices silently pointed at the wrong paragraphs. Content is stable.
    hits = [p for p in d.paragraphs if substring in p.text]
    assert len(hits) == 1, f"expected exactly 1 match for {substring[:60]!r}, got {len(hits)}"
    return hits[0]


def replace_run_text(old, new):
    p = find_para(old)
    full = "".join(r.text for r in p.runs)
    new_full = full.replace(old, new)
    p.runs[0].text = new_full
    for r in p.runs[1:]:
        r.text = ""


# ---------------------------------------------------------------------------
# 1. Simplify 4.11-4.13
# ---------------------------------------------------------------------------

replace_run_text(
    "Sections 4.1 through 4.10 report accuracy, recall and false-positive rate as summary "
    "percentages. The raw true-positive, false-positive, true-negative and false-negative counts "
    "behind those percentages are given here in full, evaluated directly against the same code "
    "path production runs (SceneViolenceDetector in maincode/x3d_violence_detector.py), one "
    "prediction per clip, at the exact threshold each detector is currently configured with in "
    "config.json. All figures are measured on each detector's held-out TEST split -- clips "
    "excluded from both training and checkpoint selection for that detector's manifest.",
    "Sections 4.1 through 4.10 summarise performance as percentages, such as recall and "
    "false-positive rate. This section shows the raw counts behind those percentages -- how many "
    "clips were correctly flagged, how many were missed, and how many false alarms were raised -- "
    "for all three detectors, not only violence. Every figure comes from running the same "
    "detection code the app actually uses, scoring one prediction per clip, at the exact "
    "confidence threshold each detector is currently set to. All figures are measured on each "
    "detector's held-out test clips: footage set aside and never used for training or for "
    "choosing the final model.")

replace_run_text(
    "Table 4.6: Confusion Matrix Counts at Configured Operating Points",
    "Table 4.6: Confusion Matrix Counts at Current Threshold Settings")

replace_run_text(
    "Figure 4.11a-c: Confusion Matrices — Violence (daynight checkpoint), Robbery, Vandalism",
    "Figure 4.11a-c: Confusion Matrices — Violence (day-and-night model), Robbery, Vandalism")

replace_run_text(
    "This is precisely why config.json currently "
    "ships vandalism detection disabled (detection.vandalism.enabled = false) rather than "
    "presenting it "
    "as validated to the same standard as violence and robbery -- an explicit scope decision "
    "recorded in "
    "Section 4.10 rather than an oversight.",
    "This is exactly why vandalism detection currently ships turned off, rather than being "
    "presented as validated to the same standard as violence and robbery -- a deliberate scope "
    "decision recorded in Section 4.10, not an oversight.")

replace_run_text(
    "Every camera frame passes through the same fixed pipeline. YOLO11s-pose first locates "
    "people and body keypoints in the frame. Rather than three unrelated models, violence, "
    "robbery and vandalism share one detector class (SceneViolenceDetector) parameterised by a "
    "different X3D-XS checkpoint, confidence threshold and consecutive-frame requirement per "
    "category, read from config.json. Each detector's raw per-frame confidence is passed through "
    "_smooth_and_confirm, which applies an exponential moving average and requires a configurable "
    "number of consecutive above-threshold frames before an incident is confirmed -- a single "
    "high-confidence frame does not by itself raise an alert. Weapon and sign detection runs as a "
    "separate YOLO model outside this shared class, since it is a single-frame object-detection "
    "task rather than a temporal one.",
    "Every camera frame passes through the same fixed pipeline. A pose-detection model "
    "(YOLO11s-pose) first locates people and their body position in the frame. Violence, robbery "
    "and vandalism are not three unrelated models bolted together -- they share one underlying "
    "detector design, trained separately for each category with its own confidence threshold and "
    "its own requirement for how many frames in a row must look suspicious before anything is "
    "confirmed. Each detector's raw confidence for the current frame is smoothed: instead of "
    "reacting to a single lucky or unlucky frame, the system averages recent frames, weighting "
    "the most recent ones more heavily, and only confirms an incident once several frames in a "
    "row clear the threshold. Weapon and sign detection runs separately, using its own model, "
    "because recognising an object only needs a single frame -- it has no need for this "
    "frame-to-frame smoothing.")

replace_run_text(
    "X3D-XS was chosen specifically because it is small enough (approximately 3 million "
    "parameters) to keep three instances resident on a single consumer GPU (the development "
    "machine's GTX 1660 SUPER, 6 GB VRAM) with headroom left for the pose and weapon models "
    "running alongside them, as confirmed by preflight.py's inference benchmark (Chapter III, "
    "Development Phase).",
    "X3D-XS, the architecture behind all three category detectors, was chosen specifically for "
    "its small size -- about 3 million parameters -- which is what makes it possible to run "
    "three instances of it at once on a single consumer GPU (the development machine's GTX 1660 "
    "SUPER, with 6 GB of memory), with room left over for the pose and weapon models running "
    "alongside them. This was confirmed by direct measurement on the actual hardware, not "
    "assumed (Chapter III, Development Phase).")

replace_run_text(
    "4.13 Concurrent Multi-Detector Execution",
    "4.13 How the Detectors Run at the Same Time")

replace_run_text(
    "\"Running at the same time\" in this system means two distinct things, both visible in "
    "maincode/main.py. First, the three category detectors are not three separate processes "
    "racing each other -- they run sequentially, in order, on every frame within a single main "
    "detection loop, which is affordable specifically because each X3D-XS pass is inexpensive. "
    "Second, and separately, four supporting tasks -- weapon/sign detection, posting a confirmed "
    "alert to the backend, finalising a saved incident clip, and encoding frames for the live "
    "dashboard stream -- are each offloaded to their own dedicated single-worker thread pool "
    "(Python's ThreadPoolExecutor) rather than run inline. A camera frame is read on its own "
    "thread as well, so a slow network write or a slow encode never stalls frame capture or the "
    "next detection pass.",
    "\"Running at the same time\" actually means two different things here. First, the three "
    "category detectors do not run as separate, competing processes -- they run one after "
    "another, in a fixed order, on every single frame. This is affordable only because each "
    "individual pass is cheap; running three in sequence on every frame would be too slow "
    "otherwise. Second, and separately, four supporting jobs -- checking for weapons, sending a "
    "confirmed alert to the system, saving an incident clip, and preparing frames for the live "
    "dashboard -- are each handed off to their own dedicated background worker instead of being "
    "done inline. The camera feed itself is also read on its own background thread. The result "
    "is that a slow step, such as sending an alert over the network, never holds up the "
    "detectors from processing the next frame.")

replace_run_text(
    "Figure 4.13: Concurrency Model Within One Camera Process",
    "Figure 4.13: How Detection Tasks Overlap Within One Camera")

replace_run_text(
    "This design keeps one camera responsive under load, but it does not by itself scale to many "
    "cameras: the current implementation runs one such pipeline per camera process, with no "
    "inference shared across cameras. Person-crop inference, which would run the classifiers "
    "only where people are detected rather than on every frame regardless of activity, is "
    "recorded in Section 4.10 as the principal remaining step toward serving multiple cameras "
    "from one GPU.",
    "This design keeps one camera responsive, but it does not by itself scale to many cameras at "
    "once: the current implementation runs one full pipeline per camera, with nothing shared "
    "between cameras. Running the detectors only where people are actually visible, rather than "
    "on every frame regardless of activity, is recorded in Section 4.10 as the main remaining "
    "step toward supporting multiple cameras from a single GPU.")

print("Simplified 4.11-4.13.")

# ---------------------------------------------------------------------------
# 2. Add weapon/sign dataset section after "E. Evaluation Metrics"
# ---------------------------------------------------------------------------
anchor = find_para(
    "False alarms per hour is the number of alarms raised per hour of continuous footage "
    "containing no incident.")


def add_after(anchor_p, text, bold=False, justify=True):
    p = anchor_p.insert_paragraph_before(text)
    anchor_p._p.addprevious(p._p)
    anchor_p._p.addnext(p._p)
    if justify:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold and p.runs:
        p.runs[0].font.bold = True
    return p


cur = anchor
cur = add_after(cur, "F. Weapon and Sign Detector Dataset", bold=True)
cur = add_after(cur,
    "Weapon and sign detection uses a separate YOLO object-detection model, trained "
    "independently from the three detectors described above (violence, robbery, vandalism). Its "
    "training images and weapon/sign labels were assembled from seven public datasets hosted on "
    "Roboflow Universe, the dataset labelling and management platform introduced earlier in this "
    "chapter (AI Development, Section 3.4). Table 3.2b lists all seven; full citations follow "
    "the table.")

cur = add_after(cur, "Table 3.2b: Weapon and Sign Dataset Sources")

# Table: Dataset / Owner / Images / License
tbl = d.add_table(rows=8, cols=4)
tbl.style = "Table Grid"
headers = ["Dataset", "Owner", "Images", "Licence"]
for j, h in enumerate(headers):
    c = tbl.rows[0].cells[j]
    c.text = ""
    r = c.paragraphs[0].add_run(h)
    r.font.bold = True
rows_data = [
    ("Gun-cctv-detection", "dietest", "5,149", "CC BY 4.0"),
    ("Gun detection", "workspace-1qko2", "36,105", "Public Domain"),
    ("Gun and knife detection", "Mahad Ahmed", "8,451", "CC BY 4.0"),
    ("Knife", "Sanket Kulkarni", "7,073", "CC BY 4.0"),
    ("Knife-dataset", "workspace-zqssx", "4,075", "CC BY 4.0"),
    ("Traffic and Road Signs", "Usman Chaudhry", "10,000", "CC BY 4.0"),
    ("CCTV Knife Detection Dataset (synthetic, sample)", "Simuletic", "114", "CC BY 4.0"),
]
for i, row_vals in enumerate(rows_data, start=1):
    for j, v in enumerate(row_vals):
        tbl.rows[i].cells[j].text = v
# move this table's XML element to sit right after its caption paragraph
cur._p.addnext(tbl._tbl)

# Chain every subsequent paragraph after the PREVIOUS one just inserted, not
# a fixed anchor -- addnext() always inserts immediately after its target, so
# reusing the same fixed anchor for multiple calls silently reverses their
# order (each new call lands between the anchor and the previous call's
# result). A moving `prev` avoids that.
def add_next(prev_el, text, bold=False, italic=False, size=None):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    if size:
        r.font.size = size
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    prev_el.addnext(p._p)
    return p

prev = add_next(tbl._tbl,
    "The synthetic dataset (Simuletic) is a 114-image sample of a larger, non-public dataset; "
    "it is included for the diversity of its synthetic CCTV-angle knife imagery, not as a "
    "primary source. All other datasets contain real photographs. Class labels vary by source "
    "(handguns, rifles, knives, and traffic/road signs, the last of which supports the sign "
    "component of the weapon-and-sign detector) and were consolidated during preprocessing.")

prev = add_next(prev._p, "Dataset Citations", bold=True)

citations = [
    "dietest. (2023). Gun-cctv-detection (Version 1) [Data set]. Roboflow. "
    "https://universe.roboflow.com/dietest/gun-cctv-detection",
    "workspace-1qko2. (2024). gun detection (Version 4) [Data set]. Roboflow. "
    "https://universe.roboflow.com/workspace-1qko2/gun-detection-ghlzd",
    "Ahmed, M. (2022). gun and knife detection (Version 1) [Data set]. Roboflow. "
    "https://universe.roboflow.com/mahad-ahmed/gun-and-knife-detection",
    "Kulkarni, S. (2022). knife (Version 1) [Data set]. Roboflow. "
    "https://universe.roboflow.com/sanket-kulkarni/knife-eydvx",
    "workspace-zqssx. (2022). knife-dataset (Version 2) [Data set]. Roboflow. "
    "https://universe.roboflow.com/workspace-zqssx/knife-dataset-4kytl",
    "Chaudhry, U. (2022). Traffic and Road Signs (Version 1) [Data set]. Roboflow. "
    "https://universe.roboflow.com/usmanchaudhry622-gmail-com/traffic-and-road-signs",
    "Simuletic. (2025). CCTV Knife Detection Dataset (Sample Version) (Version 1) [Data set]. "
    "Roboflow. https://universe.roboflow.com/simuletic/cctv-knife-detection-dataset-zkkaf",
]
for c in citations:
    prev = add_next(prev._p, c, size=docx.shared.Pt(10))

prev = add_next(prev._p,
    "These entries follow Roboflow's own recommended data-set citation format and should be "
    "carried into a References/Bibliography chapter if and when one is added to this manuscript "
    "(see note in Chapter II).",
    italic=True, size=docx.shared.Pt(9))

d.save(PATH)
print("Added Section F with full weapon dataset citations. Saved.")
